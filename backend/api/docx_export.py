"""
DOCX export functionality for EasyRead content.
Creates Word documents with text and images in an accessible format.
"""

import os
import tempfile
from io import BytesIO
from django.conf import settings
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.section import WD_SECTION
from .constants import DISCLAIMER_TEXT, GLOBALSYMBOLS_ACKNOWLEDGEMENT_TEXT
from docx.oxml.shared import OxmlElement, qn
from PIL import Image
import logging

logger = logging.getLogger(__name__)

# UNICEF Blue color (RGB: 0, 189, 242)
UNICEF_BLUE = RGBColor(0, 189, 242)
UNICEF_BLUE_HEX = "00BDF2"  # Hex representation of UNICEF blue


def add_page_border(doc, color=UNICEF_BLUE):
    """Add a border around all pages in UNICEF blue."""
    sections = doc.sections
    for section in sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        
        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')  # Border width
            border.set(qn('w:space'), '24')  # Space from edge
            border.set(qn('w:color'), UNICEF_BLUE_HEX)
            pgBorders.append(border)
        
        sectPr.append(pgBorders)


def add_page_numbers(doc):
    """Add page numbers to the document footer."""
    sections = doc.sections
    for section in sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
        # Style the page number
        run.font.size = Pt(10)
        run.font.color.rgb = UNICEF_BLUE


def create_docx_export(title, easy_read_content, original_markdown=None):
    """
    Create a DOCX document from EasyRead content.
    
    Args:
        title (str): Document title
        easy_read_content (list): List of sentence/image pairs
        original_markdown (str, optional): Original source content
    
    Returns:
        BytesIO: DOCX document as bytes
    """
    try:
        # Create a new Document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(title or 'EasyRead Document', 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_paragraph('Easy-to-Read Version')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        subtitle.runs[0].font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1
        
        # Add disclaimer and GlobalSymbols acknowledgement below subtitle
        doc.add_paragraph()  # Add some space
        
        # Add disclaimer
        disclaimer_para = doc.add_paragraph()
        disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer_run = disclaimer_para.add_run(DISCLAIMER_TEXT)
        disclaimer_run.font.size = Pt(12)
        disclaimer_run.font.italic = True
        disclaimer_run.font.color.theme_color = MSO_THEME_COLOR_INDEX.DARK_1
        
        doc.add_paragraph()  # Add space between disclaimer and acknowledgement
        
        # Add GlobalSymbols acknowledgement
        acknowledgement_para = doc.add_paragraph()
        acknowledgement_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add globe emoji and acknowledgement text
        run = acknowledgement_para.add_run("🌍 ")
        run.font.size = Pt(12)
        
        run = acknowledgement_para.add_run(GLOBALSYMBOLS_ACKNOWLEDGEMENT_TEXT)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_2
        
        # Add a page break after title
        doc.add_page_break()
        
        # Create a table for the content (2 columns)
        # Filter out empty sentences first
        valid_content = [item for item in easy_read_content if item.get('sentence', '').strip()]
        
        if valid_content:
            # Create table with 2 columns and rows for each sentence
            table = doc.add_table(rows=len(valid_content), cols=2)
            
            # Remove table borders by setting table style to None and removing borders manually
            table.style = None
            
            # Remove all table borders
            def remove_table_borders(table):
                """Remove all borders from a table"""
                tbl = table._tbl
                tblPr = tbl.tblPr
                
                # Remove table borders
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tblBorders.append(border)
                
                tblPr.append(tblBorders)
            
            remove_table_borders(table)
            
            # Set column widths (30% for image, 70% for text)
            for row in table.rows:
                row.cells[0].width = Inches(2.5)  # Image column
                row.cells[1].width = Inches(4.5)  # Text column
                
                # Add spacing between rows by setting row height
                row.height = Inches(0.3)  # Minimum row height for spacing
        
        # Process each sentence/image pair
        for idx, item in enumerate(valid_content):
            sentence = item.get('sentence', '').strip()
            image_path = item.get('selected_image_path')
            
            # Get the table row for this item
            row = table.rows[idx]
            image_cell = row.cells[0]
            text_cell = row.cells[1]
            
            # Set cell vertical alignment
            image_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            text_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Add cell margins for better spacing
            def set_cell_margins(cell, top=200, bottom=200, start=100, end=100):
                """Set cell margins in twips (1/20th of a point)"""
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                
                for margin_name, value in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), str(value))
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                
                tcPr.append(tcMar)
            
            # Apply margins to both cells for spacing
            set_cell_margins(image_cell, top=300, bottom=300)  # More vertical spacing
            set_cell_margins(text_cell, top=300, bottom=300)
            
            # Add image to the left cell
            if image_path:
                try:
                    logger.info(f"Processing image for sentence {idx}: {image_path}")
                    
                    # Handle different image path formats
                    if image_path.startswith('http://') or image_path.startswith('https://'):
                        # It's a URL - extract the relative path
                        from urllib.parse import urlparse
                        parsed_url = urlparse(image_path)
                        
                        # Extract path after /media/
                        url_path = parsed_url.path
                        if url_path.startswith('/media/'):
                            relative_path = url_path[7:]  # Remove '/media/' prefix
                        else:
                            relative_path = url_path.lstrip('/')
                        
                        logger.info(f"Extracted relative path from URL: {relative_path}")
                        
                        # Construct full file system path
                        media_root = getattr(settings, 'MEDIA_ROOT', 'media')
                        full_image_path = os.path.join(media_root, relative_path)
                        
                    elif image_path.startswith('/media/'):
                        # URL path starting with /media/ - treat like HTTP URL
                        relative_path = image_path[7:]  # Remove '/media/' prefix
                        logger.info(f"Extracted relative path from media URL: {relative_path}")
                        
                        # Construct full file system path
                        media_root = getattr(settings, 'MEDIA_ROOT', 'media')
                        full_image_path = os.path.join(media_root, relative_path)
                        
                    elif image_path.startswith('/'):
                        # Absolute file system path
                        full_image_path = image_path
                    else:
                        # Relative path - try multiple possible locations
                        media_root = getattr(settings, 'MEDIA_ROOT', 'media')
                        
                        # Remove leading slash if present
                        clean_path = image_path.lstrip('/')
                        
                        # Try different path combinations
                        possible_paths = [
                            os.path.join(media_root, clean_path),
                            os.path.join(settings.BASE_DIR.parent, 'media', clean_path),
                            os.path.join(settings.BASE_DIR.parent, clean_path),
                            os.path.join(settings.BASE_DIR, clean_path),
                            clean_path,  # Try as absolute path
                        ]
                        
                        full_image_path = None
                        for path in possible_paths:
                            if os.path.exists(path):
                                full_image_path = path
                                logger.info(f"Found image at: {full_image_path}")
                                break
                        
                        if not full_image_path:
                            # Log all attempted paths for debugging
                            logger.warning(f"Image not found at any of these paths: {possible_paths}")
                            full_image_path = possible_paths[0]  # Use first attempt for error message
                    
                    if full_image_path and os.path.exists(full_image_path):
                        # Get image dimensions to maintain aspect ratio
                        with Image.open(full_image_path) as img:
                            width, height = img.size
                            logger.info(f"Image dimensions: {width}x{height}")
                            
                            # Calculate scaled dimensions for 150px height (approximately 1.04 inches)
                            target_height_inches = 1.04  # 150px at 144 DPI
                            aspect_ratio = width / height
                            scaled_height = target_height_inches
                            scaled_width = aspect_ratio * target_height_inches
                            
                            # Ensure width doesn't exceed cell width (2.5 inches)
                            max_width = 2.2  # Leave some padding
                            if scaled_width > max_width:
                                scaled_width = max_width
                                scaled_height = scaled_width / aspect_ratio
                        
                        # Add image to the image cell
                        img_para = image_cell.paragraphs[0]
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_para.add_run()
                        run.add_picture(full_image_path, width=Inches(scaled_width), height=Inches(scaled_height))
                        logger.info(f"Successfully added image to table cell: {full_image_path}")
                        
                    else:
                        logger.warning(f"Image not found: {full_image_path}")
                        # Add placeholder text if image not found
                        img_para = image_cell.paragraphs[0]
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_para.add_run(f"[Image not found]")
                        run.font.size = Pt(10)
                        run.font.italic = True
                        run.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_2
                        
                except Exception as e:
                    logger.error(f"Error adding image {image_path}: {str(e)}")
                    # Add placeholder text if image fails
                    img_para = image_cell.paragraphs[0]
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_para.add_run(f"[Image error]")
                    run.font.size = Pt(10)
                    run.font.italic = True
                    run.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_2
            else:
                # No image - add placeholder
                img_para = image_cell.paragraphs[0]
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run("[No image]")
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_2
            
            # Add sentence text to the right cell
            text_para = text_cell.paragraphs[0]
            text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add extra spacing between lines
            text_para.paragraph_format.space_after = Pt(12)  # Space after paragraph
            text_para.paragraph_format.line_spacing = Pt(18)  # Line height
            
            run = text_para.add_run(sentence)
            run.font.size = Pt(14)  # Larger font for accessibility
        
        # Add original content section if provided
        if original_markdown:
            doc.add_page_break()
            doc.add_heading('Original Content', level=1)
            
            # Add original markdown in a formatted way
            original_para = doc.add_paragraph()
            run = original_para.add_run(original_markdown)
            run.font.name = 'Courier New'  # Monospace font
            run.font.size = Pt(10)
        
        # Apply page styling
        add_page_border(doc)
        add_page_numbers(doc)
        
        # Save to BytesIO
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        return docx_buffer
        
    except Exception as e:
        logger.error(f"Error creating DOCX export: {str(e)}")
        raise


def get_safe_filename(title):
    """
    Convert a title to a safe filename for DOCX export.
    
    Args:
        title (str): Original title
        
    Returns:
        str: Safe filename without extension
    """
    if not title:
        return "easyread_document"
    
    # Remove/replace unsafe characters
    safe_chars = []
    for char in title:
        if char.isalnum() or char in ' -_':
            safe_chars.append(char)
        else:
            safe_chars.append('_')
    
    safe_title = ''.join(safe_chars).strip()
    
    # Limit length and remove extra spaces
    safe_title = ' '.join(safe_title.split())[:50]
    
    return safe_title.lower().replace(' ', '_') if safe_title else "easyread_document"